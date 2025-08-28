import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter
import PyPDF2


def load_mapping(mapping_file: Path) -> dict:
    """Load word mappings from a file."""
    mapping = {}
    with mapping_file.open('r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or '=' not in line:
                continue
            src, dst = line.split('=', 1)
            mapping[src] = dst
    return mapping


def convert_pdf_to_docx_preserve_structure(pdf_path: Path, docx_path: Path):
    """Convert PDF to DOCX while preserving structure and formatting."""
    try:
        # Use pdf2docx to convert with structure preservation
        cv = Converter(str(pdf_path))
        cv.convert(str(docx_path), start=0, end=None)
        cv.close()
        print(f"Converted PDF to DOCX with structure preserved: {docx_path}")
        return True
    except Exception as e:
        print(f"Error converting PDF to DOCX: {e}")
        return False


def replace_in_paragraph(paragraph, mapping):
    """Replace text in a paragraph using the mapping."""
    for src, dst in mapping.items():
        if src in paragraph.text:
            for run in paragraph.runs:
                if src in run.text:
                    run.text = run.text.replace(src, dst)
    
    # Also check if the paragraph itself contains text that needs replacement
    for src, dst in mapping.items():
        if src in paragraph.text:
            # Rebuild the paragraph text with replacements
            new_text = paragraph.text
            for src, dst in mapping.items():
                new_text = new_text.replace(src, dst)
            
            # Clear existing runs and add new text
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                # Remove additional runs if any
                for run in paragraph.runs[1:]:
                    run.text = ""


def replace_in_table(table, mapping):
    """Replace text in table cells."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, mapping)


def replace_in_header_footer(header_footer, mapping):
    """Replace text in headers and footers."""
    # Process all paragraphs in the header/footer
    for paragraph in header_footer.paragraphs:
        replace_in_paragraph(paragraph, mapping)
    
    # Process all tables in the header/footer
    for table in header_footer.tables:
        replace_in_table(table, mapping)
    
    # Process any shapes or text boxes in the header/footer
    if hasattr(header_footer, 'shapes'):
        for shape in header_footer.shapes:
            if hasattr(shape, 'text_frame'):
                for paragraph in shape.text_frame.paragraphs:
                    replace_in_paragraph(paragraph, mapping)


def translate_document_preserve_formatting(doc_path: Path, mapping: dict, output_dir: Path):
    """Translate a DOCX document while preserving formatting."""
    doc = Document(doc_path)
    
    # Process main document paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, mapping)
    
    # Process main document tables
    for table in doc.tables:
        replace_in_table(table, mapping)
    
    # Process any shapes or text boxes in the main document
    if hasattr(doc, 'shapes'):
        for shape in doc.shapes:
            if hasattr(shape, 'text_frame'):
                for paragraph in shape.text_frame.paragraphs:
                    replace_in_paragraph(paragraph, mapping)
    
    # Process all sections including headers and footers
    for section in doc.sections:
        replace_in_header_footer(section.header, mapping)
        replace_in_header_footer(section.footer, mapping)
    
    # Direct XML text replacement using a more aggressive approach
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Process all text elements in the document XML
        for element in doc._element.iter():
            # Process text content
            if hasattr(element, 'text') and element.text:
                original_text = element.text
                modified = False
                for src, dst in mapping.items():
                    if src in original_text:
                        original_text = original_text.replace(src, dst)
                        modified = True
                
                if modified:
                    # Use a more direct approach to modify text
                    try:
                        # Clear the text first
                        element.text = ""
                        # Set the new text
                        element.text = original_text
                    except Exception as e:
                        print(f"Could not modify text for element {element.tag}: {e}")
            
            # Process tail text (text after elements)
            if hasattr(element, 'tail') and element.tail:
                original_tail = element.tail
                modified = False
                for src, dst in mapping.items():
                    if src in original_tail:
                        original_tail = original_tail.replace(src, dst)
                        modified = True
                
                if modified:
                    try:
                        element.tail = original_tail
                    except Exception as e:
                        print(f"Could not modify tail text for element {element.tag}: {e}")
                        
    except Exception as e:
        print(f"Warning: Could not process XML elements: {e}")
    
    # Final pass: force update all paragraphs
    try:
        for paragraph in doc.paragraphs:
            if paragraph.text:
                original_text = paragraph.text
                modified = False
                for src, dst in mapping.items():
                    if src in original_text:
                        original_text = original_text.replace(src, dst)
                        modified = True
                
                if modified:
                    # Clear all runs and create a new one with the translated text
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = original_text
                    else:
                        paragraph.add_run(original_text)
    except Exception as e:
        print(f"Warning: Could not process final paragraph pass: {e}")
    
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{doc_path.stem}_translated{doc_path.suffix}"
    doc.save(output_path)
    print(f"Translated document saved: {output_path}")
    return output_path


def process_pdf_to_structured_docx(pdf_path: Path, mapping_file: Path, output_dir: Path):
    """Main function to process PDF to structured DOCX with translation."""
    print(f"Processing PDF: {pdf_path}")
    
    # Step 1: Convert PDF to DOCX with structure preservation
    print("Step 1: Converting PDF to DOCX with structure preservation...")
    temp_docx_path = output_dir / f"{pdf_path.stem}_structured.docx"
    
    if not convert_pdf_to_docx_preserve_structure(pdf_path, temp_docx_path):
        print("Failed to convert PDF to DOCX with structure preservation")
        return None
    
    # Step 2: Load translation mapping
    print("Step 2: Loading translation mapping...")
    mapping = load_mapping(mapping_file)
    print(f"Loaded {len(mapping)} translation mappings")
    
    # Step 3: Translate the DOCX while preserving formatting
    print("Step 3: Translating document while preserving formatting...")
    translated_path = translate_document_preserve_formatting(temp_docx_path, mapping, output_dir)
    
    # Step 4: Clean up temporary file
    try:
        temp_docx_path.unlink()
        print("Cleaned up temporary file")
    except:
        pass
    
    return translated_path


def main():
    input_dir = Path('invoice')
    output_dir = Path('translated')
    mapping_file = Path('words.txt')
    
    if not input_dir.exists():
        raise FileNotFoundError(f"Input directory {input_dir} not found")
    
    if not mapping_file.exists():
        raise FileNotFoundError(f"Mapping file {mapping_file} not found")
    
    # Process all PDF files in the invoice directory
    pdf_files = list(input_dir.glob('*.pdf'))
    if not pdf_files:
        print("No PDF files found in the invoice directory")
        return
    
    print(f"Found {len(pdf_files)} PDF file(s) to process")
    
    for pdf_path in pdf_files:
        print(f"\n{'='*50}")
        print(f"Processing: {pdf_path.name}")
        print(f"{'='*50}")
        
        try:
            result = process_pdf_to_structured_docx(pdf_path, mapping_file, output_dir)
            if result:
                print(f"Successfully processed: {pdf_path.name} -> {result.name}")
            else:
                print(f"Failed to process: {pdf_path.name}")
        except Exception as e:
            print(f"Error processing {pdf_path.name}: {e}")


if __name__ == '__main__':
    main()

