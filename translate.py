import os
from pathlib import Path
from docx import Document


def load_mapping(mapping_file: Path) -> dict:
    """Load word mappings from a file."""
    mapping = {}
    with mapping_file.open('r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or '=' not in line:
                continue
            src, dst = line.split('=', 1)
            mapping[src] = dst
    return mapping


def replace_in_paragraph(paragraph, mapping):
    for src, dst in mapping.items():
        if src in paragraph.text:
            for run in paragraph.runs:
                if src in run.text:
                    run.text = run.text.replace(src, dst)


def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, mapping)


def replace_in_header_footer(header_footer, mapping):
    for paragraph in header_footer.paragraphs:
        replace_in_paragraph(paragraph, mapping)
    for table in header_footer.tables:
        replace_in_table(table, mapping)


def translate_document(doc_path: Path, mapping: dict, output_dir: Path):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, mapping)
    for table in doc.tables:
        replace_in_table(table, mapping)
    for section in doc.sections:
        replace_in_header_footer(section.header, mapping)
        replace_in_header_footer(section.footer, mapping)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{doc_path.stem}_translated{doc_path.suffix}"
    doc.save(output_path)
    return output_path


def main():
    input_dir = Path('invoice')
    output_dir = Path('translated')
    mapping_file = Path('words.txt')

    if not input_dir.exists():
        raise FileNotFoundError(f"Input directory {input_dir} not found")

    mapping = load_mapping(mapping_file)
    for doc_path in input_dir.glob('*.docx'):
        translate_document(doc_path, mapping, output_dir)


if __name__ == '__main__':
    main()
